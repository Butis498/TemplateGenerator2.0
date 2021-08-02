import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pathlib import Path

class TempleteGenerator():

    def __init__(self):
        self.required_files = ['Replacement_Dictionary.csv','db.csv']
        DB_folder_name = 'DB'
        cwd = os.getcwd() 
        directory = str(os.path.join(cwd,DB_folder_name))
        if not os.path.isdir(directory):
            os.mkdir(directory)

    
        try:   
            self.validate_files_required(self.required_files,directory)
        except FileExistsError as e:

            for file in self.required_files:
                open(os.path.join(directory,file),'w')


        self.df =pd.read_csv(os.path.join(directory,'Replacement_Dictionary.csv'))

    
        self.json_data = {}
        for _, row in self.df.iterrows():
            self.json_data[row['REPLACE_KEYWORD']] = row['REPLACEMENT']

        self.Keydictionary = list(self.json_data.keys())



        self.db =pd.read_csv(os.path.join(directory,'db.csv')) 
        self.file = open(os.path.join(directory,"db.csv"))
        self.data_count = len(self.file.readlines()) - 1
        self.file.close()
        self.Generated_folder_name = "\\Generated"
        self.path = str(os.getcwd() + self.Generated_folder_name)


        if not os.path.isdir(os.getcwd() + self.Generated_folder_name):
            os.mkdir(os.getcwd() + self.Generated_folder_name)


    def replacedParograph(self , parograph , key , data_to_replace):
        try:
            replaced = parograph.replace(key , data_to_replace)
        except TypeError as _:
            replaced = ""

        return replaced

    def validate_files_required(self,file_list,directory):

        for file in file_list:
            if file not in os.listdir(directory):
                raise FileExistsError('Missing files for operation')


    
    def get_Data(self , i , key):
        return str(self.db[self.json_data[key]].iloc[i])

    def isfloat(self , value):
        try:
            float(value)
            return True
        except ValueError:
            return False

    def replaceContentXls(self ,documentName ):

        for i in range(self.data_count):
            
            document = load_workbook(documentName)

            for sheet in document._sheets:

                for row in range(1,sheet.max_row + 1):
                    for col in range(1,sheet.max_column + 1):
                        par = str(sheet.cell(row,col).value)
                        if par == 'None':
                            par = ""
                            
                        for key in self.Keydictionary:
                            if key in par:
                                data_to_replace = self.get_Data(i , key)
                                par = self.replacedParograph(par , key , data_to_replace )
                            else:
                                if self.isfloat(par):
                                    float_value = float(par)
                                    sheet.cell(row, col).value = float_value
                                else:
                                    sheet.cell(row, col).value = par

                        if self.isfloat(par):
                            float_value = float(par)
                            sheet.cell(row, col).value = float_value
                        else:
                            sheet.cell(row, col).value = par

            path, file = os.path.split(documentName)
            file_name , file_extension = os.path.splitext(file)
            path = Path(path)
            path = str(path.parent)
            document.save(self.path  + "\\"+ file_name + "_" + str(i) + file_extension)
        


    def replaceContentDoc(self , documentName ):

        for i in range(self.data_count):
            document = Document(documentName)

            for paragraph in document.paragraphs:
                for key in self.Keydictionary:
                    if key in paragraph.text:
                        data_to_replace = self.get_Data(i , key)
                        paragraph.text = self.replacedParograph(paragraph.text , key , data_to_replace)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key in self.Keydictionary:
                                if key in paragraph.text:
                                    data_to_replace = self.get_Data(i , key)
                                    paragraph.text = self.replacedParograph(paragraph.text , key , data_to_replace)
            path, file = os.path.split(documentName)
            file_name , file_extension = os.path.splitext(file)
            path = Path(path)
            path = str(path.parent)
            document.save(self.path  + "\\"+ file_name + "_" + str(i) + file_extension)
                    
    


